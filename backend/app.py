from flask import Flask, request, jsonify
from flask_cors import CORS
from werkzeug.security import generate_password_hash, check_password_hash
from pymongo import MongoClient
from bson import ObjectId
from datetime import datetime, timedelta
import secrets
import os
from functools import wraps
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)
CORS(app)
app.config['SECRET_KEY'] = os.getenv('SECRET_KEY', 'your-secret-key-here')

# MongoDB Atlas connection
MONGO_URI = os.getenv("MONGO_URI", "mongodb+srv://hari:hari919597@cluster1.p2aip1m.mongodb.net/?retryWrites=true&w=majority&appName=Cluster1")
client = MongoClient(MONGO_URI)
db = client.subscription_manager

# Collections
users_collection = db.users
plans_collection = db.subscription_plans
subscriptions_collection = db.user_subscriptions
discounts_collection = db.discounts

# Simple token storage (in production, use Redis or database)
active_tokens = {}

def generate_token():
    return secrets.token_urlsafe(32)

def token_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization')
        print(f"Received token: {token}")
        
        if not token:
            print("No token provided")
            return jsonify({'error': 'Token is missing'}), 401
        
        try:
            # Handle both 'Bearer token' and 'token' formats
            if token.startswith('Bearer '):
                token = token.split(' ')[1]
            
            print(f"Processed token: {token}")
            print(f"Active tokens: {list(active_tokens.keys())}")
            
            if token not in active_tokens:
                print(f"Token not found in active tokens")
                return jsonify({'error': 'Token is invalid or expired'}), 401
            
            user_id = active_tokens[token]
            print(f"Found user_id: {user_id}")
            
            current_user = users_collection.find_one({'_id': ObjectId(user_id)})
            if not current_user:
                print(f"User not found for id: {user_id}")
                return jsonify({'error': 'User not found'}), 401
                
            print(f"User found: {current_user['email']}")
        except Exception as e:
            print(f"Token validation error: {e}")
            return jsonify({'error': 'Token is invalid'}), 401
        return f(current_user, *args, **kwargs)
    return decorated

@app.route('/register', methods=['POST'])
def register():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    name = data.get('name', email.split('@')[0])
    role = data.get('role', 'user')
    
    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400
    
    if users_collection.find_one({'email': email}):
        return jsonify({'error': 'Email already exists'}), 409
    
    hashed_password = generate_password_hash(password)
    user_data = {
        'email': email,
        'password': hashed_password,
        'name': name,
        'role': role,
        'createdAt': datetime.utcnow(),
        'updatedAt': datetime.utcnow()
    }
    
    result = users_collection.insert_one(user_data)
    return jsonify({'message': 'User registered successfully', 'userId': str(result.inserted_id)}), 201

@app.route('/admin/register', methods=['POST'])
def admin_register():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    name = data.get('name', email.split('@')[0])
    
    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400
    
    if users_collection.find_one({'email': email}):
        return jsonify({'error': 'Email already exists'}), 409
    
    hashed_password = generate_password_hash(password)
    user_data = {
        'email': email,
        'password': hashed_password,
        'name': name,
        'role': 'admin',
        'createdAt': datetime.utcnow(),
        'updatedAt': datetime.utcnow()
    }
    
    result = users_collection.insert_one(user_data)
    return jsonify({'message': 'Admin registered successfully', 'userId': str(result.inserted_id)}), 201

@app.route('/login', methods=['POST'])
def login():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    
    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400
    
    user = users_collection.find_one({'email': email})
    
    if user and check_password_hash(user['password'], password):
        token = generate_token()
        active_tokens[token] = str(user['_id'])
        
        return jsonify({
            'message': 'Login successful',
            'token': token,
            'user': {
                'id': str(user['_id']),
                'email': user['email'],
                'name': user['name'],
                'role': user['role']
            }
        }), 200
    else:
        return jsonify({'error': 'Invalid credentials'}), 401

@app.route('/admin/login', methods=['POST'])
def admin_login():
    data = request.get_json()
    email = data.get('email')
    password = data.get('password')
    
    if not email or not password:
        return jsonify({'error': 'Email and password required'}), 400
    
    user = users_collection.find_one({'email': email, 'role': 'admin'})
    
    if user and check_password_hash(user['password'], password):
        token = generate_token()
        active_tokens[token] = str(user['_id'])
        
        return jsonify({
            'message': 'Admin login successful',
            'token': token,
            'user': {
                'id': str(user['_id']),
                'email': user['email'],
                'name': user['name'],
                'role': user['role']
            }
        }), 200
    else:
        return jsonify({'error': 'Invalid admin credentials'}), 401

@app.route('/logout', methods=['POST'])
@token_required
def logout(current_user):
    token = request.headers.get('Authorization').split(' ')[1]
    if token in active_tokens:
        del active_tokens[token]
    return jsonify({'message': 'Logged out successfully'}), 200

@app.route('/plans', methods=['GET'])
def get_plans():
    plans = list(plans_collection.find())
    for plan in plans:
        plan['_id'] = str(plan['_id'])
    return jsonify(plans), 200

@app.route('/plans', methods=['POST'])
@token_required
def create_plan(current_user):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    data = request.get_json()
    plan_data = {
        'name': data.get('name'),
        'description': data.get('description'),
        'price': data.get('price'),
        'interval': data.get('interval', 'monthly'),
        'features': data.get('features', []),
        'isPopular': data.get('isPopular', False),
        'maxUsers': data.get('maxUsers'),
        'storage': data.get('storage'),
        'support': data.get('support'),
        'createdAt': datetime.utcnow(),
        'updatedAt': datetime.utcnow()
    }
    
    result = plans_collection.insert_one(plan_data)
    plan_data['_id'] = str(result.inserted_id)
    return jsonify(plan_data), 201

@app.route('/plans/<plan_id>', methods=['PUT'])
@token_required
def update_plan(current_user, plan_id):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    data = request.get_json()
    update_data = {
        'name': data.get('name'),
        'description': data.get('description'),
        'price': data.get('price'),
        'interval': data.get('interval', 'monthly'),
        'features': data.get('features', []),
        'isPopular': data.get('isPopular', False),
        'maxUsers': data.get('maxUsers'),
        'storage': data.get('storage'),
        'support': data.get('support'),
        'updatedAt': datetime.utcnow()
    }
    
    result = plans_collection.update_one(
        {'_id': ObjectId(plan_id)},
        {'$set': update_data}
    )
    
    if result.matched_count == 0:
        return jsonify({'error': 'Plan not found'}), 404
    
    updated_plan = plans_collection.find_one({'_id': ObjectId(plan_id)})
    updated_plan['_id'] = str(updated_plan['_id'])
    return jsonify(updated_plan), 200

@app.route('/plans/<plan_id>', methods=['DELETE'])
@token_required
def delete_plan(current_user, plan_id):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    # Check if plan exists
    plan = plans_collection.find_one({'_id': ObjectId(plan_id)})
    if not plan:
        return jsonify({'error': 'Plan not found'}), 404
    
    # Cancel all active subscriptions for this plan
    subscriptions_collection.update_many(
        {'planId': plan_id, 'status': 'active'},
        {'$set': {
            'status': 'cancelled',
            'autoRenew': False,
            'updatedAt': datetime.utcnow()
        }}
    )
    
    # Delete the plan
    result = plans_collection.delete_one({'_id': ObjectId(plan_id)})
    
    if result.deleted_count == 0:
        return jsonify({'error': 'Failed to delete plan'}), 500
    
    return jsonify({'message': 'Plan deleted successfully'}), 200

@app.route('/subscriptions', methods=['GET'])
@token_required
def get_user_subscriptions(current_user):
    subscriptions = list(subscriptions_collection.find({'userId': str(current_user['_id'])}))
    for sub in subscriptions:
        sub['_id'] = str(sub['_id'])
        # Get plan details
        plan = plans_collection.find_one({'_id': ObjectId(sub['planId'])})
        if plan:
            plan['_id'] = str(plan['_id'])
            sub['plan'] = plan
    return jsonify(subscriptions), 200

@app.route('/subscriptions', methods=['POST'])
@token_required
def create_subscription(current_user):
    data = request.get_json()
    plan_id = data.get('planId')
    discount_code = data.get('discountCode', '').upper() if data.get('discountCode') else None
    
    if not plan_id:
        return jsonify({'error': 'Plan ID required'}), 400
    
    # Cancel any existing active subscriptions (allow plan changes)
    subscriptions_collection.update_many(
        {
            'userId': str(current_user['_id']),
            'status': 'active'
        },
        {
            '$set': {
                'status': 'cancelled',
                'autoRenew': False,
                'updatedAt': datetime.utcnow()
            }
        }
    )
    
    plan = plans_collection.find_one({'_id': ObjectId(plan_id)})
    if not plan:
        return jsonify({'error': 'Plan not found'}), 404
    
    original_price = plan['price']
    final_price = original_price
    discount_applied = None
    
    # Handle discount code if provided
    if discount_code:
        discount = discounts_collection.find_one({'code': discount_code})
        
        if discount and discount.get('isActive', True):
            now = datetime.utcnow()
            
            # Check timing and usage limits
            if (not discount.get('startDate') or now >= discount['startDate']) and \
               (not discount.get('endDate') or now <= discount['endDate']) and \
               (not discount.get('usageLimit') or discount.get('usedCount', 0) < discount['usageLimit']) and \
               (not discount.get('applicablePlans') or plan_id in discount['applicablePlans']):
                
                # Calculate discount
                if discount['type'] == 'percentage':
                    discount_amount = (original_price * discount['value']) / 100
                else:  # fixed
                    discount_amount = min(discount['value'], original_price)
                
                final_price = max(0, original_price - discount_amount)
                discount_applied = {
                    'code': discount['code'],
                    'type': discount['type'],
                    'value': discount['value'],
                    'discountAmount': discount_amount
                }
                
                # Increment usage count
                discounts_collection.update_one(
                    {'_id': discount['_id']},
                    {'$inc': {'usedCount': 1}}
                )
    
    subscription_data = {
        'userId': str(current_user['_id']),
        'planId': plan_id,
        'status': 'active',
        'startDate': datetime.utcnow(),
        'endDate': (datetime.utcnow() + timedelta(days=30 if plan['interval'] == 'monthly' else 365)),
        'autoRenew': data.get('autoRenew', True),
        'paymentMethod': data.get('paymentMethod'),
        'originalPrice': original_price,
        'finalPrice': final_price,
        'discountApplied': discount_applied,
        'createdAt': datetime.utcnow(),
        'updatedAt': datetime.utcnow()
    }
    
    result = subscriptions_collection.insert_one(subscription_data)
    subscription_data['_id'] = str(result.inserted_id)
    return jsonify(subscription_data), 201

@app.route('/subscriptions/<subscription_id>/cancel', methods=['PUT'])
@token_required
def cancel_subscription(current_user, subscription_id):
    subscription = subscriptions_collection.find_one({
        '_id': ObjectId(subscription_id),
        'userId': str(current_user['_id'])
    })
    
    if not subscription:
        return jsonify({'error': 'Subscription not found'}), 404
    
    if subscription['status'] != 'active':
        return jsonify({'error': 'Subscription is not active'}), 400
    
    # Update subscription status to cancelled
    result = subscriptions_collection.update_one(
        {'_id': ObjectId(subscription_id)},
        {'$set': {
            'status': 'cancelled',
            'autoRenew': False,
            'updatedAt': datetime.utcnow()
        }}
    )
    
    if result.matched_count == 0:
        return jsonify({'error': 'Failed to cancel subscription'}), 500
    
    return jsonify({'message': 'Subscription cancelled successfully'}), 200

@app.route('/subscriptions/<subscription_id>/payment', methods=['PUT'])
@token_required
def update_payment_method(current_user, subscription_id):
    data = request.get_json()
    payment_method = data.get('paymentMethod')
    
    if not payment_method:
        return jsonify({'error': 'Payment method required'}), 400
    
    subscription = subscriptions_collection.find_one({
        '_id': ObjectId(subscription_id),
        'userId': str(current_user['_id'])
    })
    
    if not subscription:
        return jsonify({'error': 'Subscription not found'}), 404
    
    result = subscriptions_collection.update_one(
        {'_id': ObjectId(subscription_id)},
        {'$set': {
            'paymentMethod': payment_method,
            'updatedAt': datetime.utcnow()
        }}
    )
    
    if result.matched_count == 0:
        return jsonify({'error': 'Failed to update payment method'}), 500
    
    return jsonify({'message': 'Payment method updated successfully'}), 200

@app.route('/subscriptions/<subscription_id>/invoice', methods=['GET'])
@token_required
def download_invoice(current_user, subscription_id):
    from flask import make_response
    import io
    
    try:
        subscription = subscriptions_collection.find_one({
            '_id': ObjectId(subscription_id),
            'userId': str(current_user['_id'])
        })
        
        if not subscription:
            return jsonify({'error': 'Subscription not found'}), 404
        
        plan = plans_collection.find_one({'_id': ObjectId(subscription['planId'])})
        if not plan:
            return jsonify({'error': 'Plan not found'}), 404
        
        # Try Word document generation first
        try:
            from docx import Document
            from docx.shared import Inches
            
            doc = Document()
            
            # Title
            title = doc.add_heading('INVOICE', 0)
            title.alignment = 1  # Center alignment
            
            # Invoice info
            doc.add_paragraph(f"Invoice #: INV-{subscription_id[:8]}")
            doc.add_paragraph(f"Date: {datetime.utcnow().strftime('%Y-%m-%d')}")
            doc.add_paragraph()
            
            # Customer info
            doc.add_heading('Bill To:', level=2)
            doc.add_paragraph(f"Customer: {current_user['name']}")
            doc.add_paragraph(f"Email: {current_user['email']}")
            doc.add_paragraph()
            
            # Subscription details
            doc.add_heading('Subscription Details:', level=2)
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Description'
            hdr_cells[1].text = 'Period'
            hdr_cells[2].text = 'Amount'
            
            # Data row
            row_cells = table.add_row().cells
            row_cells[0].text = f"{plan['name']} - {plan.get('description', 'Subscription Plan')}"
            row_cells[1].text = f"{subscription['startDate'].strftime('%Y-%m-%d')} to {subscription['endDate'].strftime('%Y-%m-%d')}"
            row_cells[2].text = f"₹{subscription.get('originalPrice', plan['price'])}"
            
            # Discount row if applicable
            if subscription.get('discountApplied'):
                discount_row = table.add_row().cells
                discount_row[0].text = f"Discount ({subscription['discountApplied']['code']})"
                discount_row[1].text = "-"
                discount_row[2].text = f"-₹{subscription['discountApplied']['discountAmount']:.2f}"
            
            doc.add_paragraph()
            
            # Total
            total_para = doc.add_paragraph()
            total_para.add_run(f"Total Amount: ₹{subscription.get('finalPrice', plan['price'])}").bold = True
            
            doc.add_paragraph()
            doc.add_paragraph("Thank you for your business!")
            
            # Save to memory
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            
            response = make_response(doc_io.getvalue())
            response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            response.headers['Content-Disposition'] = f'attachment; filename=invoice-{subscription_id[:8]}.docx'
            
            return response
            
        except ImportError:
            # Fallback to simple text format
            invoice_content = f"""INVOICE

Invoice #: INV-{subscription_id[:8]}
Date: {datetime.utcnow().strftime('%Y-%m-%d')}

Bill To:
Customer: {current_user['name']}
Email: {current_user['email']}

Subscription Details:
Plan: {plan['name']}
Description: {plan.get('description', 'N/A')}
Price: ₹{subscription.get('finalPrice', plan['price'])}/{plan['interval']}
Start Date: {subscription['startDate'].strftime('%Y-%m-%d')}
End Date: {subscription['endDate'].strftime('%Y-%m-%d')}
Status: {subscription['status']}

{f"Discount Applied: {subscription['discountApplied']['code']} (-₹{subscription['discountApplied']['discountAmount']:.2f})" if subscription.get('discountApplied') else ""}

Total Amount: ₹{subscription.get('finalPrice', plan['price'])}

Thank you for your business!"""
            
            response = make_response(invoice_content)
            response.headers['Content-Type'] = 'text/plain'
            response.headers['Content-Disposition'] = f'attachment; filename=invoice-{subscription_id[:8]}.txt'
            
            return response
            
    except Exception as e:
        print(f"Invoice generation error: {e}")
        return jsonify({'error': 'Failed to generate invoice'}), 500

@app.route('/admin/users', methods=['GET'])
@token_required
def get_all_users(current_user):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    # Get all users
    users = list(users_collection.find({'role': 'user'}))
    
    result = []
    for user in users:
        user_id_str = str(user['_id'])
        
        # Get active subscriptions for this user
        active_subs = list(subscriptions_collection.find({
            'userId': user_id_str,
            'status': 'active'
        }))
        
        user_data = {
            'id': user_id_str,
            'name': user['name'],
            'email': user['email'],
            'createdAt': user['createdAt'].isoformat() if user.get('createdAt') else None,
            'activeSubscriptions': []
        }
        
        for sub in active_subs:
            plan = plans_collection.find_one({'_id': ObjectId(sub['planId'])})
            if plan:
                user_data['activeSubscriptions'].append({
                    'planName': plan['name'],
                    'planPrice': plan['price'],
                    'startDate': sub['startDate'].isoformat() if sub.get('startDate') else None,
                    'endDate': sub['endDate'].isoformat() if sub.get('endDate') else None
                })
        
        result.append(user_data)
    
    return jsonify(result), 200

@app.route('/analytics', methods=['GET'])
@token_required
def get_analytics(current_user):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    # Calculate analytics
    total_subscriptions = subscriptions_collection.count_documents({'status': 'active'})
    
    # Calculate total revenue
    pipeline = [
        {'$match': {'status': 'active'}},
        {'$lookup': {
            'from': 'subscription_plans',
            'localField': 'planId',
            'foreignField': '_id',
            'as': 'plan'
        }},
        {'$unwind': '$plan'},
        {'$group': {
            '_id': None,
            'totalRevenue': {'$sum': '$plan.price'}
        }}
    ]
    
    revenue_result = list(subscriptions_collection.aggregate(pipeline))
    total_revenue = revenue_result[0]['totalRevenue'] if revenue_result else 0
    
    # Top plans
    top_plans_pipeline = [
        {'$match': {'status': 'active'}},
        {'$group': {
            '_id': '$planId',
            'subscriptions': {'$sum': 1}
        }},
        {'$lookup': {
            'from': 'subscription_plans',
            'localField': '_id',
            'foreignField': '_id',
            'as': 'plan'
        }},
        {'$unwind': '$plan'},
        {'$project': {
            'planId': {'$toString': '$_id'},
            'planName': '$plan.name',
            'subscriptions': 1,
            'revenue': {'$multiply': ['$subscriptions', '$plan.price']}
        }},
        {'$sort': {'subscriptions': -1}},
        {'$limit': 5}
    ]
    
    top_plans = list(subscriptions_collection.aggregate(top_plans_pipeline))
    
    analytics = {
        'totalRevenue': total_revenue,
        'activeSubscriptions': total_subscriptions,
        'churnRate': 5.2,  # Mock data
        'topPlans': top_plans,
        'monthlyGrowth': 12.5  # Mock data
    }
    
    return jsonify(analytics), 200

@app.route('/recommendations/<user_id>', methods=['GET'])
@token_required
def get_recommendations(current_user, user_id):
    # Simple AI recommendation logic
    user_subscriptions = list(subscriptions_collection.find({'userId': user_id, 'status': 'active'}))
    all_plans = list(plans_collection.find())
    
    recommendations = []
    for plan in all_plans[:3]:  # Limit to 3 recommendations
        plan['_id'] = str(plan['_id'])
        recommendation = {
            'planId': plan['_id'],
            'planName': plan['name'],
            'reason': f"Based on your usage patterns, {plan['name']} offers better value",
            'confidence': 85,
            'potentialSavings': 50,
            'features': plan['features']
        }
        recommendations.append(recommendation)
    
    return jsonify(recommendations), 200

# Discount Management Routes
@app.route('/admin/discounts', methods=['GET'])
@token_required
def get_discounts(current_user):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    discounts = list(discounts_collection.find())
    for discount in discounts:
        discount['_id'] = str(discount['_id'])
    return jsonify(discounts), 200

@app.route('/admin/discounts', methods=['POST'])
@token_required
def create_discount(current_user):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    data = request.get_json()
    
    # Validate required fields
    if not data.get('code') or not data.get('type') or not data.get('value'):
        return jsonify({'error': 'Code, type, and value are required'}), 400
    
    # Check if discount code already exists
    if discounts_collection.find_one({'code': data.get('code')}):
        return jsonify({'error': 'Discount code already exists'}), 409
    
    discount_data = {
        'code': data.get('code').upper(),
        'type': data.get('type'),  # 'percentage' or 'fixed'
        'value': data.get('value'),
        'description': data.get('description', ''),
        'startDate': datetime.fromisoformat(data.get('startDate')) if data.get('startDate') else datetime.utcnow(),
        'endDate': datetime.fromisoformat(data.get('endDate')) if data.get('endDate') else None,
        'usageLimit': data.get('usageLimit'),
        'usedCount': 0,
        'isActive': data.get('isActive', True),
        'applicablePlans': data.get('applicablePlans', []),
        'createdAt': datetime.utcnow(),
        'updatedAt': datetime.utcnow()
    }
    
    result = discounts_collection.insert_one(discount_data)
    discount_data['_id'] = str(result.inserted_id)
    return jsonify(discount_data), 201

@app.route('/admin/discounts/<discount_id>', methods=['DELETE'])
@token_required
def delete_discount(current_user, discount_id):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    result = discounts_collection.delete_one({'_id': ObjectId(discount_id)})
    
    if result.deleted_count == 0:
        return jsonify({'error': 'Discount not found'}), 404
    
    return jsonify({'message': 'Discount deleted successfully'}), 200

@app.route('/admin/discounts/<discount_id>', methods=['PUT'])
@token_required
def update_discount(current_user, discount_id):
    if current_user['role'] != 'admin':
        return jsonify({'error': 'Admin access required'}), 403
    
    data = request.get_json()
    
    update_data = {
        'description': data.get('description'),
        'startDate': datetime.fromisoformat(data.get('startDate')) if data.get('startDate') else None,
        'endDate': datetime.fromisoformat(data.get('endDate')) if data.get('endDate') else None,
        'usageLimit': data.get('usageLimit'),
        'isActive': data.get('isActive'),
        'applicablePlans': data.get('applicablePlans', []),
        'updatedAt': datetime.utcnow()
    }
    
    # Remove None values
    update_data = {k: v for k, v in update_data.items() if v is not None}
    
    result = discounts_collection.update_one(
        {'_id': ObjectId(discount_id)},
        {'$set': update_data}
    )
    
    if result.matched_count == 0:
        return jsonify({'error': 'Discount not found'}), 404
    
    updated_discount = discounts_collection.find_one({'_id': ObjectId(discount_id)})
    updated_discount['_id'] = str(updated_discount['_id'])
    return jsonify(updated_discount), 200

@app.route('/discounts/active', methods=['GET'])
def get_active_discounts():
    now = datetime.utcnow()
    
    # Find active discounts that are currently valid
    active_discounts = list(discounts_collection.find({
        'isActive': True,
        '$or': [
            {'startDate': {'$lte': now}},
            {'startDate': None}
        ],
        '$or': [
            {'endDate': {'$gte': now}},
            {'endDate': None}
        ]
    }))
    
    result = []
    for discount in active_discounts:
        # Check if usage limit is reached
        if discount.get('usageLimit') and discount.get('usedCount', 0) >= discount['usageLimit']:
            continue
            
        result.append({
            'code': discount['code'],
            'type': discount['type'],
            'value': discount['value'],
            'description': discount['description'],
            'endDate': discount.get('endDate').isoformat() if discount.get('endDate') else None
        })
    
    return jsonify(result), 200

@app.route('/discounts/validate', methods=['POST'])
@token_required
def validate_discount(current_user):
    data = request.get_json()
    code = data.get('code', '').upper()
    plan_id = data.get('planId')
    
    if not code:
        return jsonify({'error': 'Discount code required'}), 400
    
    discount = discounts_collection.find_one({'code': code})
    
    if not discount:
        return jsonify({'error': 'Invalid discount code'}), 404
    
    # Check if discount is active
    if not discount.get('isActive', True):
        return jsonify({'error': 'Discount code is inactive'}), 400
    
    # Check timing
    now = datetime.utcnow()
    if discount.get('startDate') and now < discount['startDate']:
        return jsonify({'error': 'Discount code is not yet active'}), 400
    
    if discount.get('endDate') and now > discount['endDate']:
        return jsonify({'error': 'Discount code has expired'}), 400
    
    # Check usage limit
    if discount.get('usageLimit') and discount.get('usedCount', 0) >= discount['usageLimit']:
        return jsonify({'error': 'Discount code usage limit reached'}), 400
    
    # Check applicable plans
    if discount.get('applicablePlans') and plan_id not in discount['applicablePlans']:
        return jsonify({'error': 'Discount code not applicable to this plan'}), 400
    
    # Calculate discount amount
    if plan_id:
        plan = plans_collection.find_one({'_id': ObjectId(plan_id)})
        if plan:
            if discount['type'] == 'percentage':
                discount_amount = (plan['price'] * discount['value']) / 100
            else:  # fixed
                discount_amount = min(discount['value'], plan['price'])
            
            final_price = max(0, plan['price'] - discount_amount)
            
            return jsonify({
                'valid': True,
                'discount': {
                    'id': str(discount['_id']),
                    'code': discount['code'],
                    'type': discount['type'],
                    'value': discount['value'],
                    'discountAmount': discount_amount,
                    'originalPrice': plan['price'],
                    'finalPrice': final_price
                }
            }), 200
    
    return jsonify({
        'valid': True,
        'discount': {
            'id': str(discount['_id']),
            'code': discount['code'],
            'type': discount['type'],
            'value': discount['value']
        }
    }), 200

if __name__ == '__main__':
    app.run(debug=True, port=5000)